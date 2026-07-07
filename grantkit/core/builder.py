"""Assemble a grant into a single document (``grantkit build``).

Given a :class:`~grantkit.core.project.GrantProject` this module produces:

* a compiled document in ``md`` / ``html`` / ``pdf`` / ``docx`` — for
  plain-text portals (``accepts_markdown: false``) the compiled text is a set
  of labelled copy blocks ready to paste into portal boxes;
* optionally (``--share``) a single self-contained ``assembled.html`` review
  page with per-section word counts and status badges;
* always ``status.json``.

PDF needs the ``pdf`` extra (WeasyPrint); DOCX needs the ``docx`` extra
(python-docx). Missing extras raise :class:`BuildDependencyError` with an
actionable message rather than failing opaquely.
"""

from __future__ import annotations

import html as _html
from dataclasses import dataclass, field
from pathlib import Path
from typing import TYPE_CHECKING, Optional

import markdown as _markdown

from ..utils.text import count_words
from .checks import CheckResult
from .status import build_status, days_until_deadline, write_status

if TYPE_CHECKING:  # pragma: no cover - typing only
    from .project import GrantProject, SectionState

VALID_FORMATS = ("md", "html", "pdf", "docx")

_STATUS_COLORS = {
    "complete": ("#137333", "#e6f4ea"),
    "partial": ("#8a6d00", "#fef7e0"),
    "empty": ("#5f6368", "#f1f3f4"),
    "over_limit": ("#c5221f", "#fce8e6"),
}


class BuildDependencyError(Exception):
    """Raised when an optional format dependency is not installed."""


@dataclass
class BuildResult:
    """Paths written by a build."""

    format: str
    document_path: Optional[Path] = None
    share_path: Optional[Path] = None
    status_path: Optional[Path] = None
    warnings: list[str] = field(default_factory=list)

    def outputs(self) -> list[Path]:
        return [
            p
            for p in (
                self.document_path,
                self.share_path,
                self.status_path,
            )
            if p is not None
        ]


def build_project(
    project: "GrantProject",
    *,
    fmt: str = "md",
    share: bool = False,
    output: Optional[Path] = None,
    checks: Optional[CheckResult] = None,
) -> BuildResult:
    """Compile ``project`` and always refresh ``status.json``."""
    if fmt not in VALID_FORMATS:
        raise ValueError(
            f"Unknown format '{fmt}'. Choose from {', '.join(VALID_FORMATS)}."
        )

    result = BuildResult(format=fmt)

    ext = fmt
    doc_path = output or (project.root / f"proposal.{ext}")
    doc_path.parent.mkdir(parents=True, exist_ok=True)

    if fmt == "md":
        doc_path.write_text(_compile_text(project), encoding="utf-8")
    elif fmt == "html":
        doc_path.write_text(_compile_document_html(project), encoding="utf-8")
    elif fmt == "pdf":
        _write_pdf(project, doc_path)
    elif fmt == "docx":
        _write_docx(project, doc_path)
    result.document_path = doc_path

    if share:
        share_path = project.root / "assembled.html"
        share_path.write_text(_share_page(project, checks), encoding="utf-8")
        result.share_path = share_path

    result.status_path = write_status(project, checks)
    return result


# -- text assembly ------------------------------------------------------


def _compile_text(project: "GrantProject") -> str:
    """The compiled document body as text (Markdown or plain-text blocks)."""
    if project.accepts_markdown:
        return _assemble_markdown(project)
    return _assemble_plaintext_blocks(project)


def _assemble_markdown(project: "GrantProject") -> str:
    parts: list[str] = []
    title = project.title or project.funder or "Grant proposal"
    parts.append(f"# {title}\n")
    if project.program:
        parts.append(f"*{project.program}*\n")
    for section in project.sections:
        parts.append(f"\n## {section.title}\n")
        body = section.body.strip() if section.exists else ""
        parts.append(body if body else "_(empty)_")
        parts.append("")
    return "\n".join(parts).strip() + "\n"


def _assemble_plaintext_blocks(project: "GrantProject") -> str:
    """Plain-text copy blocks for portals that accept plain text only."""
    parts: list[str] = []
    header = project.title or project.funder or "Grant proposal"
    parts.append(header)
    parts.append("=" * len(header))
    parts.append(
        "Plain-text portal: paste each block below into its portal box.\n"
    )
    for section in project.sections:
        if section.format == "fields":
            # Individual form fields, typed one by one — keep the source
            # (tables and all) as an on-screen reference, not a paste block.
            label = f"{section.title}  (form fields — enter individually)"
            body = section.body.strip() if section.exists else ""
        else:
            body = _to_plaintext(section.body) if section.exists else ""
            limit = f" / {section.word_limit}" if section.word_limit else ""
            label = f"{section.title}  ({section.words}{limit} words)"
        parts.append("-" * 70)
        parts.append(label)
        parts.append("-" * 70)
        parts.append(body if body else "(empty)")
        parts.append("")
    return "\n".join(parts).strip() + "\n"


def _to_plaintext(markdown_text: str) -> str:
    """Best-effort Markdown -> plain text for portal copy blocks."""
    import re

    text = markdown_text
    text = re.sub(r"<!--.*?-->", "", text, flags=re.DOTALL)  # HTML comments
    text = re.sub(r"^#{1,6}\s+", "", text, flags=re.MULTILINE)  # headers
    text = re.sub(r"\[([^\]]+)\]\([^\)]+\)", r"\1", text)  # links
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)  # bold
    text = re.sub(r"__([^_]+)__", r"\1", text)  # bold
    text = re.sub(r"(?<!\*)\*([^*]+)\*(?!\*)", r"\1", text)  # italic
    text = re.sub(r"`([^`]+)`", r"\1", text)  # inline code
    text = re.sub(r"^\s*>\s?", "", text, flags=re.MULTILINE)  # quotes
    return text.strip()


# -- html / pdf / docx --------------------------------------------------


def _compile_document_html(project: "GrantProject") -> str:
    if project.accepts_markdown:
        body_html = _markdown.markdown(
            _assemble_markdown(project),
            extensions=["tables", "fenced_code"],
        )
    else:
        blocks = _assemble_plaintext_blocks(project)
        body_html = f"<pre>{_html.escape(blocks)}</pre>"
    title = _html.escape(project.title or project.funder or "Grant proposal")
    return _document_html_shell(title, body_html)


def _document_html_shell(title: str, body_html: str) -> str:
    return (
        '<!doctype html>\n<html lang="en">\n<head>\n'
        '<meta charset="utf-8">\n'
        '<meta name="viewport" content="width=device-width, '
        'initial-scale=1">\n'
        f"<title>{title}</title>\n"
        "<style>\n"
        "body{font-family:Georgia,'Times New Roman',serif;max-width:46rem;"
        "margin:2rem auto;padding:0 1rem;line-height:1.5;color:#202124}\n"
        "h1,h2,h3{font-family:-apple-system,Segoe UI,Roboto,sans-serif}\n"
        "pre{white-space:pre-wrap;font-family:inherit}\n"
        "</style>\n</head>\n<body>\n"
        f"{body_html}\n</body>\n</html>\n"
    )


def _write_pdf(project: "GrantProject", out_path: Path) -> None:
    try:
        from weasyprint import HTML
    except Exception as exc:  # ImportError or missing system libs
        raise BuildDependencyError(
            "PDF output needs WeasyPrint. Install it with "
            "`pip install grantkit[pdf]` (WeasyPrint also needs the system "
            f"Pango/Cairo libraries). Underlying error: {exc}"
        )
    html_doc = _compile_document_html(project)
    HTML(string=html_doc).write_pdf(str(out_path))


def _write_docx(project: "GrantProject", out_path: Path) -> None:
    try:
        from docx import Document
    except Exception as exc:
        raise BuildDependencyError(
            "DOCX output needs python-docx. Install it with "
            f"`pip install grantkit[docx]`. Underlying error: {exc}"
        )
    document = Document()
    document.add_heading(
        project.title or project.funder or "Grant proposal", level=0
    )
    if project.program:
        document.add_paragraph(project.program)
    for section in project.sections:
        document.add_heading(section.title, level=1)
        body = _to_plaintext(section.body) if section.exists else "(empty)"
        for para in body.split("\n\n"):
            document.add_paragraph(para.strip())
    document.save(str(out_path))


# -- share review page --------------------------------------------------


def _share_page(project: "GrantProject", checks: Optional[CheckResult]) -> str:
    status = build_status(project, checks)
    grant = status["grant"]
    completion = status["completion"]
    check_block = status["checks"]

    countdown = days_until_deadline(project.deadline)
    if countdown is None:
        deadline_txt = grant["deadline"] or "no deadline set"
    elif countdown < 0:
        deadline_txt = f"{grant['deadline']} ({abs(countdown)} days ago)"
    else:
        deadline_txt = f"{grant['deadline']} (in {countdown} days)"

    cards = "\n".join(_section_card(section) for section in project.sections)
    title = _html.escape(grant["title"] or grant["funder"] or "Grant proposal")
    subtitle = _html.escape(
        " · ".join(x for x in (grant["funder"], grant["program"]) if x)
    )
    errors = check_block["errors"]
    warnings = check_block["warnings"]
    check_summary = _check_summary_html(check_block)

    return f"""<!doctype html>
<html lang="en">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>{title} — GrantKit review</title>
<style>
:root{{color-scheme:light dark}}
body{{font-family:-apple-system,Segoe UI,Roboto,Helvetica,Arial,sans-serif;
margin:0;background:#f8f9fa;color:#202124}}
.wrap{{max-width:52rem;margin:0 auto;padding:2rem 1rem 4rem}}
header h1{{margin:0 0 .25rem;font-size:1.6rem}}
header p{{margin:.1rem 0;color:#5f6368}}
.meta{{display:flex;flex-wrap:wrap;gap:1rem;margin:1.25rem 0}}
.stat{{background:#fff;border:1px solid #dadce0;border-radius:10px;
padding:.75rem 1rem;flex:1;min-width:8rem}}
.stat .n{{font-size:1.5rem;font-weight:600}}
.stat .l{{color:#5f6368;font-size:.8rem}}
.card{{background:#fff;border:1px solid #dadce0;border-radius:10px;
padding:1rem 1.15rem;margin:.6rem 0;display:flex;justify-content:space-between;
align-items:center;gap:1rem}}
.card .t{{font-weight:600}}
.card .w{{color:#5f6368;font-size:.85rem;margin-top:.15rem}}
.badge{{padding:.2rem .6rem;border-radius:999px;font-size:.75rem;
font-weight:600;white-space:nowrap}}
.checks{{margin-top:1.5rem}}
.checks li{{margin:.2rem 0}}
.err{{color:#c5221f}}.warn{{color:#8a6d00}}
footer{{margin-top:2rem;color:#5f6368;font-size:.8rem}}
@media (prefers-color-scheme:dark){{
body{{background:#202124;color:#e8eaed}}
.stat,.card{{background:#2d2e30;border-color:#5f6368}}
header p,.stat .l,.card .w,footer{{color:#9aa0a6}}}}
</style>
</head>
<body>
<div class="wrap">
<header>
<h1>{title}</h1>
<p>{subtitle}</p>
<p>Deadline: {_html.escape(deadline_txt)}</p>
</header>
<div class="meta">
<div class="stat"><div class="n">{completion['percent']:.0f}%</div>
<div class="l">complete</div></div>
<div class="stat"><div class="n">{completion['sections_complete']}/{completion['sections_total']}</div>
<div class="l">sections done</div></div>
<div class="stat"><div class="n">{completion['words_total']:,}</div>
<div class="l">words</div></div>
<div class="stat"><div class="n {('err' if errors else '')}">{errors}</div>
<div class="l">errors</div></div>
<div class="stat"><div class="n {('warn' if warnings else '')}">{warnings}</div>
<div class="l">warnings</div></div>
</div>
{cards}
{check_summary}
<footer>Generated by GrantKit {status['grantkit_version']} · {status['generated_at']}</footer>
</div>
</body>
</html>
"""


def _section_card(section: "SectionState") -> str:
    color, bg = _STATUS_COLORS.get(section.status, ("#5f6368", "#f1f3f4"))
    limit = f" / {section.word_limit}" if section.word_limit else ""
    words = f"{section.words}{limit} words"
    label = section.status.replace("_", " ")
    return (
        '<div class="card"><div><div class="t">'
        f"{_html.escape(section.title)}</div>"
        f'<div class="w">{words}</div></div>'
        f'<span class="badge" style="color:{color};background:{bg}">'
        f"{label}</span></div>"
    )


def _check_summary_html(check_block: dict) -> str:
    items = check_block["items"]
    if not items:
        return (
            '<div class="checks"><h2>Checks</h2>'
            "<p>No issues found.</p></div>"
        )
    lines = []
    for item in items:
        cls = "err" if item["level"] == "error" else "warn"
        where = f" [{item['section']}]" if item.get("section") else ""
        lines.append(
            f'<li class="{cls}"><strong>{item["level"]}</strong>'
            f"{_html.escape(where)}: {_html.escape(item['message'])}</li>"
        )
    return (
        '<div class="checks"><h2>Checks</h2><ul>'
        + "\n".join(lines)
        + "</ul></div>"
    )


def word_count(text: str) -> int:
    """Convenience re-export for callers/tests."""
    return count_words(text)
